# backend/document_processor.py
import os
from typing import List
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

# Try to load docx support
try:
    from langchain_community.document_loaders import Docx2txtLoader
    DOCX_LOADER_AVAILABLE = True
except ImportError:
    DOCX_LOADER_AVAILABLE = False
    print("⚠️ Docx2txtLoader not available. Install: pip install docx2txt")

class DocumentProcessor:
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""],
            length_function=len
        )
        print("✅ DocumentProcessor ready")
    
    def load_document(self, file_path: str) -> List[Document]:
        ext = os.path.splitext(file_path)[1].lower()
        print(f"📄 Loading: {os.path.basename(file_path)}")
        
        if ext == ".pdf":
            loader = PyPDFLoader(file_path)
            docs = loader.load()
        
        elif ext in [".txt", ".md"]:
            loader = TextLoader(file_path, encoding="utf-8")
            docs = loader.load()
        
        elif ext == ".docx":
            if DOCX_LOADER_AVAILABLE:
                from langchain_community.document_loaders import Docx2txtLoader
                loader = Docx2txtLoader(file_path)
                docs = loader.load()
            else:
                try:
                    import docx
                    doc = docx.Document(file_path)
                    full_text = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            full_text.append(paragraph.text)
                    content = "\n".join(full_text)
                    docs = [Document(page_content=content, metadata={"source": os.path.basename(file_path)})]
                except ImportError:
                    raise ImportError("Please install docx2txt: pip install docx2txt")
        
        else:
            raise ValueError(f"Unsupported file type: {ext}")
        
        for d in docs:
            d.metadata["source"] = os.path.basename(file_path)
            d.metadata["source_path"] = file_path
        
        print(f"   ✅ Loaded {len(docs)} sections")
        return docs
    
    def split_documents(self, documents: List[Document]) -> List[Document]:
        chunks = self.text_splitter.split_documents(documents)
        print(f"   ✅ Split into {len(chunks)} chunks")
        return chunks
    
    def process_file(self, file_path: str) -> List[Document]:
        docs = self.load_document(file_path)
        chunks = self.split_documents(docs)
        return chunks